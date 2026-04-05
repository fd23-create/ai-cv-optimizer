import fitz  # PyMuPDF
import pdfplumber
import docx
import io
import logging
from typing import Optional, Dict, Any

logger = logging.getLogger(__name__)

class PDFParser:
    """Parser for PDF and DOCX files"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.doc', '.docx']
    
    def extract_text(self, file) -> Optional[str]:
        """
        Extract text from uploaded file
        
        Args:
            file: File object from Flask request
            
        Returns:
            Extracted text as string or None if failed
        """
        if not file or not file.filename:
            return None
        
        file_extension = self._get_file_extension(file.filename)
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(file)
        elif file_extension in ['.doc', '.docx']:
            return self._extract_from_docx(file)
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return None
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename"""
        return filename.lower().split('.')[-1] if '.' in filename else ''
    
    def _extract_from_pdf(self, file) -> Optional[str]:
        """Extract text from PDF using PyMuPDF as primary, pdfplumber as fallback"""
        try:
            # Reset file pointer
            file.seek(0)
            file_bytes = file.read()
            
            # Try PyMuPDF first (faster)
            text = self._extract_with_pymupdf(file_bytes)
            if text and len(text.strip()) > 100:
                return text
            
            # Fallback to pdfplumber (better for complex layouts)
            text = self._extract_with_pdfplumber(file_bytes)
            if text and len(text.strip()) > 100:
                return text
            
            logger.warning("Both PDF extraction methods returned minimal text")
            return text or ""
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return None
    
    def _extract_with_pymupdf(self, file_bytes: bytes) -> str:
        """Extract text using PyMuPDF"""
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                
                # Basic text quality check
                if len(page_text.strip()) > 10:
                    text += page_text + "\n"
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            logger.error(f"PyMuPDF extraction failed: {str(e)}")
            return ""
    
    def _extract_with_pdfplumber(self, file_bytes: bytes) -> str:
        """Extract text using pdfplumber"""
        try:
            text = ""
            
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text and len(page_text.strip()) > 10:
                        text += page_text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"pdfplumber extraction failed: {str(e)}")
            return ""
    
    def _extract_from_docx(self, file) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            # Reset file pointer
            file.seek(0)
            
            # Create a Document object
            doc = docx.Document(file)
            
            # Extract text from paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return None
    
    def get_metadata(self, file) -> Optional[Dict[str, Any]]:
        """
        Extract metadata from file
        
        Args:
            file: File object from Flask request
            
        Returns:
            Dictionary with metadata or None if failed
        """
        if not file or not file.filename:
            return None
        
        file_extension = self._get_file_extension(file.filename)
        
        if file_extension == '.pdf':
            return self._get_pdf_metadata(file)
        elif file_extension in ['.doc', '.docx']:
            return self._get_docx_metadata(file)
        else:
            return None
    
    def _get_pdf_metadata(self, file) -> Optional[Dict[str, Any]]:
        """Extract metadata from PDF"""
        try:
            file.seek(0)
            file_bytes = file.read()
            
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            metadata = doc.metadata
            
            result = {
                'title': metadata.get('title', ''),
                'author': metadata.get('author', ''),
                'subject': metadata.get('subject', ''),
                'creator': metadata.get('creator', ''),
                'producer': metadata.get('producer', ''),
                'creationDate': metadata.get('creationDate', ''),
                'modificationDate': metadata.get('modDate', ''),
                'pageCount': len(doc),
                'format': 'PDF'
            }
            
            doc.close()
            return result
            
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {str(e)}")
            return None
    
    def _get_docx_metadata(self, file) -> Optional[Dict[str, Any]]:
        """Extract metadata from DOCX"""
        try:
            file.seek(0)
            doc = docx.Document(file)
            core_props = doc.core_properties
            
            result = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'paragraphCount': len(doc.paragraphs),
                'tableCount': len(doc.tables),
                'format': 'DOCX'
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {str(e)}")
            return None
    
    def validate_file_structure(self, text: str) -> Dict[str, Any]:
        """
        Validate the structure of extracted text
        
        Args:
            text: Extracted text
            
        Returns:
            Dictionary with validation results
        """
        if not text:
            return {
                'valid': False,
                'issues': ['No text extracted']
            }
        
        issues = []
        warnings = []
        
        # Check minimum length
        if len(text) < 100:
            issues.append('Text too short - may be incomplete')
        elif len(text) < 500:
            warnings.append('Text relatively short - verify completeness')
        
        # Check for common resume sections
        common_sections = [
            'experience', 'education', 'skills', 'summary', 'objective',
            'expériences', 'formation', 'compétences', 'résumé', 'objectif'
        ]
        
        text_lower = text.lower()
        found_sections = [section for section in common_sections if section in text_lower]
        
        if len(found_sections) < 2:
            warnings.append('Few standard resume sections detected')
        
        # Check for contact information patterns
        email_patterns = ['@', '.com', '.fr', '.org']
        phone_patterns = ['+', '(', ')', '-']
        
        has_email = any(pattern in text_lower for pattern in email_patterns)
        has_phone = any(pattern in text for pattern in phone_patterns)
        
        if not has_email:
            warnings.append('No email address detected')
        
        if not has_phone:
            warnings.append('No phone number detected')
        
        # Check for date patterns (experience, education)
        import re
        date_pattern = r'\b(19|20)\d{2}\b'
        dates = re.findall(date_pattern, text)
        
        if len(dates) < 2:
            warnings.append('Few dates detected - may lack experience/education timeline')
        
        return {
            'valid': len(issues) == 0,
            'issues': issues,
            'warnings': warnings,
            'found_sections': found_sections,
            'has_contact_info': has_email or has_phone,
            'date_count': len(dates),
            'word_count': len(text.split()),
            'char_count': len(text)
        }
